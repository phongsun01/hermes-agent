import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from ooxml_helpers import parse_xml, nsdecls, qn

def format_blocks(docx_path, format_spec):
    """
    Định dạng khối (Block Level):
    - Set margin cho các section (nếu chưa set).
    - Căn lề, spacing cho đoạn văn (đặc biệt là list items).
    - Định dạng viền, padding cho bảng biểu (trừ bảng header NĐ30).
    - Căn giữa ảnh minh họa.
    """
    doc = Document(str(docx_path))
    
    doc_type = format_spec.get("doc_type", "van_ban_dai")
    line_spacing = format_spec.get("line_spacing", 1.5)
    indent_pt = format_spec.get("first_line_indent_pt", 28)
    
    if doc_type == "hanh_chinh_nd30":
        # NĐ 30: margin đặc thù cho trang dọc
        for section in doc.sections:
            if section.orientation != WD_ORIENT.LANDSCAPE:
                from docx.shared import Mm
                section.top_margin = Mm(20)
                section.bottom_margin = Mm(20)
                section.left_margin = Mm(30)
                section.right_margin = Mm(20)
                section.page_width = Mm(210)
                section.page_height = Mm(297)

    # === PARAGRAPHS ===
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        text = para.text.strip()

        # Headings: NĐ30 không có màu, left align
        if doc_type == "hanh_chinh_nd30" and style_name.startswith('Heading'):
            para.paragraph_format.first_line_indent = Pt(indent_pt)
            
            # Xóa dấu : ở cuối (nếu có)
            if text.endswith(':'):
                for run in reversed(para.runs):
                    if run.text and run.text.strip():
                        run.text = run.text.rstrip()
                        if run.text.endswith(':'):
                            run.text = run.text[:-1]
                        break
            if style_name == 'Heading 1':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.page_break_before = False
            elif style_name == 'Heading 2' or style_name == 'Heading 3':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(3)
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    for old in pPr.findall(qn('w:pBdr')):
                        pPr.remove(old)
            continue

        # Ảnh
        p_elem = para._element
        has_drawing = bool(p_elem.findall('.//' + qn('w:drawing')))
        if has_drawing:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            continue

        # Dòng trống
        if not text:
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue

        # List items: đổi bullet thành gạch ngang (NĐ 30)
        pPr_elem = para._element.find(qn('w:pPr'))
        has_numPr = pPr_elem is not None and pPr_elem.find(qn('w:numPr')) is not None
        is_list = style_name.startswith('List') or style_name in ('Compact', 'List Paragraph') or has_numPr
        
        if is_list:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.first_line_indent = Pt(indent_pt)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.line_spacing = line_spacing
            
            if doc_type == "hanh_chinh_nd30":
                for run in para.runs:
                    if run.text:
                        run.text = run.text.replace('\u2022', '-').replace('\u25cf', '-').replace('\u25cb', '-')
                
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        pPr.remove(numPr)
                        first_text = para.text.strip()
                        if first_text and not first_text.startswith('-') and not first_text.startswith('–'):
                            if para.runs:
                                para.runs[0].text = '- ' + para.runs[0].text
            continue

        # Body text (NĐ 30):
        if doc_type == "hanh_chinh_nd30":
            if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                para.paragraph_format.first_line_indent = Cm(0)
                continue
            
            if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                para.paragraph_format.first_line_indent = Cm(0)
                continue
                
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(indent_pt)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = line_spacing

    # === TABLES ===
    if doc_type == "hanh_chinh_nd30":
        for table in doc.tables:
            tbl = table._tbl
            
            # Skip bảng header (border=nil)
            tblPr_check = tbl.find(qn('w:tblPr'))
            if tblPr_check is not None:
                tblBorders = tblPr_check.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    top_border = tblBorders.find(qn('w:top'))
                    if top_border is not None and top_border.get(qn('w:val')) == 'nil':
                        continue
            
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                tbl.insert(0, tblPr)

            for old in tblPr.findall(qn('w:tblW')):
                tblPr.remove(old)
            tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))

            for old in tblPr.findall(qn('w:tblLayout')):
                tblPr.remove(old)
            tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>'))

            for old in tblPr.findall(qn('w:tblBorders')):
                tblPr.remove(old)
            tblPr.append(parse_xml(f'''<w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tblBorders>'''))

            for old in tblPr.findall(qn('w:tblCellMar')):
                tblPr.remove(old)
            tblPr.append(parse_xml(f'''<w:tblCellMar {nsdecls("w")}>
                <w:top w:w="40" w:type="dxa"/>
                <w:left w:w="80" w:type="dxa"/>
                <w:bottom w:w="40" w:type="dxa"/>
                <w:right w:w="80" w:type="dxa"/>
            </w:tblCellMar>'''))

            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                        tc.insert(0, tcPr)

                    for old in tcPr.findall(qn('w:shd')):
                        tcPr.remove(old)

                    for para in cell.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0)
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)
                        para.paragraph_format.line_spacing = 1.15
                        
                        if i == 0:
                            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Center images explicitly
    for para in doc.paragraphs:
        p_elem = para._element
        if p_elem.findall('.//' + qn('w:drawing')):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                p_elem.insert(0, pPr)
            for old in pPr.findall(qn('w:jc')):
                pPr.remove(old)
            pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
            for old in pPr.findall(qn('w:ind')):
                pPr.remove(old)
            pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="0" w:left="0"/>'))

    doc.save(str(docx_path))

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python 03_block.py <input_docx> <format_spec.json>")
        sys.exit(1)
        
    input_docx = Path(sys.argv[1])
    format_spec_path = Path(sys.argv[2])
    
    with open(format_spec_path, "r", encoding="utf-8") as f:
        format_spec = json.load(f)
        
    print("[INFO] Bước 3: Định dạng Block (Paragraph, Table)...")
    format_blocks(input_docx, format_spec)
    print("[OK] Xong Layer 3.")
