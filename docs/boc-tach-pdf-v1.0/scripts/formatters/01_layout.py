import sys
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from ooxml_helpers import parse_xml, nsdecls, qn

def apply_landscape_sections(docx_path, format_spec):
    """
    Tìm marker [PAGE_MARKER_N], nếu N nằm trong nhóm landscape thì 
    áp dụng section break landscape.
    """
    orientations = format_spec.get("page_orientations", {})
    if not orientations:
        print("[INFO] Không có trang landscape. Bỏ qua.")
        return
    
    # Chuyển key từ string (JSON) sang int
    orientations = {int(k): v for k, v in orientations.items()}
    landscape_pages = sorted(orientations.keys())
    print(f"[INFO] Xử lý landscape cho trang: {landscape_pages}")

    doc = Document(str(docx_path))
    
    page_markers = {}  # {page_num: paragraph_index}
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.match(r"\[PAGE_MARKER_(\d+)\]", text)
        if match:
            page_markers[int(match.group(1))] = idx
            
    print(f"[INFO] Tìm thấy {len(page_markers)} page markers trong DOCX")
    
    if not page_markers:
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith("PAGE_MARKER_"):
                try:
                    num = int(text.replace("PAGE_MARKER_", "").strip("[]"))
                    page_markers[num] = idx
                except ValueError:
                    continue
                    
    if not page_markers:
        print("[WARN] Không tìm thấy page markers. Áp dụng fallback...")
        max_page = max(landscape_pages)
        total_pages = len(format_spec.get("images", []))
        if max_page >= total_pages - 2:
            last_section = doc.sections[-1]
            last_section.orientation = WD_ORIENT.LANDSCAPE
            last_section.page_width = Mm(297)
            last_section.page_height = Mm(210)
            print(f"[OK] Section cuối → landscape (fallback)")
        doc.save(str(docx_path))
        return
        
    # Nhóm landscape pages liên tục
    groups = []
    current_group = [landscape_pages[0]]
    for i in range(1, len(landscape_pages)):
        if landscape_pages[i] == landscape_pages[i-1] + 1:
            current_group.append(landscape_pages[i])
        else:
            groups.append(current_group)
            current_group = [landscape_pages[i]]
    groups.append(current_group)
    print(f"[INFO] Nhóm landscape: {groups}")
    
    for group in groups:
        first_page = group[0]
        last_page = group[-1]
        
        if first_page in page_markers:
            marker_idx = page_markers[first_page]
            
            if marker_idx < len(doc.paragraphs) and marker_idx > 0:
                prev_para = doc.paragraphs[marker_idx - 1]
                pPr = prev_para._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                    prev_para._element.insert(0, pPr)
                
                for old in pPr.findall(qn('w:sectPr')):
                    pPr.remove(old)
                
                # PORTRAIT -> Kết thúc khối dọc
                sectPr = parse_xml(f'''<w:sectPr {nsdecls("w")}>
                    <w:pgSz w:w="11906" w:h="16838"/>
                    <w:pgMar w:top="1134" w:right="1134" w:bottom="1134" 
                             w:left="1701" w:header="454" w:footer="454"/>
                </w:sectPr>''')
                pPr.append(sectPr)
                print(f"[OK] Section break portrait KẾT THÚC trước trang {first_page}")
                
        next_page = last_page + 1
        if next_page in page_markers:
            next_marker_idx = page_markers[next_page]
            if next_marker_idx > 0:
                prev_para = doc.paragraphs[next_marker_idx - 1]
                pPr = prev_para._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                    prev_para._element.insert(0, pPr)
                
                for old in pPr.findall(qn('w:sectPr')):
                    pPr.remove(old)
                
                # LANDSCAPE -> Kết thúc khối ngang
                sectPr = parse_xml(f'''<w:sectPr {nsdecls("w")}>
                    <w:pgSz w:w="16838" w:h="11906" w:orient="landscape"/>
                    <w:pgMar w:top="1134" w:right="1701" w:bottom="1134" 
                             w:left="1134" w:header="454" w:footer="454"/>
                </w:sectPr>''')
                pPr.append(sectPr)
                print(f"[OK] Section break landscape KẾT THÚC trước trang {next_page}")
        else:
            last_section = doc.sections[-1]
            last_section.orientation = WD_ORIENT.LANDSCAPE
            last_section.page_width = Mm(297)
            last_section.page_height = Mm(210)
            print(f"[OK] Section cuối → landscape")
            
    # Xóa markers và số trang thừa
    markers_to_remove = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.match(r"\[?PAGE_MARKER_(\d+)\]?", text)
        if match:
            markers_to_remove.append(para)
            page_num = int(match.group(1))
            
            # Kiểm tra xem đoạn không rỗng tiếp theo có phải là số trang tương ứng hay không
            next_idx = idx + 1
            while next_idx < len(doc.paragraphs) and not doc.paragraphs[next_idx].text.strip():
                next_idx += 1
            if next_idx < len(doc.paragraphs):
                next_para = doc.paragraphs[next_idx]
                next_text = next_para.text.strip()
                if next_text == str(page_num):
                    markers_to_remove.append(next_para)
            
    for para in markers_to_remove:
        parent = para._element.getparent()
        if parent is not None:
            parent.remove(para._element)
            
    print(f"[INFO] Đã xóa {len(markers_to_remove)} page markers và số trang thừa")
    doc.save(str(docx_path))

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python 01_layout.py <input_docx> <format_spec.json>")
        sys.exit(1)
        
    input_docx = Path(sys.argv[1])
    format_spec_path = Path(sys.argv[2])
    
    with open(format_spec_path, "r", encoding="utf-8") as f:
        format_spec = json.load(f)
        
    print("[INFO] Bước 1: Áp dụng layout trang (Orientation, Page Markers)...")
    apply_landscape_sections(input_docx, format_spec)
    print("[OK] Xong Layer 1.")
