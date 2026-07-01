import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ooxml_helpers import (
    make_run, 
    set_cell_no_border, 
    add_separator_line, 
    apply_tcW, 
    set_table_fixed_width, 
    apply_table_grid,
    remove_table_borders
)

def build_nd30_header(docx_path):
    """
    Tái cấu trúc header NĐ 30: bảng 2 cột không viền.
    Xóa các paragraph chứa thông tin cũ và thay bằng bảng header chuẩn.
    """
    doc = Document(str(docx_path))
    body_font = "Times New Roman"
    
    header_info = {
        "sao_y": None,
        "co_quan_line1": None,
        "co_quan_line2": None,
        "quoc_hieu": None,
        "tieu_ngu": None,
        "so_ky_hieu": None,
        "ngay_thang": None,
        "ten_van_ban": None,
        "trich_yeu": None,
        "kinh_gui": None,
    }
    
    header_end_idx = 0
    
    for idx in range(min(25, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        
        if not text:
            continue
            
        # Split by double spaces to separate merged columns
        segments = [s.strip() for s in re.split(r'\s{2,}', text) if s.strip()]
        
        for segment in segments:
            segment_upper = segment.upper()
            
            # 1. SAO Y
            if segment.startswith("SAO Y") and idx < 3:
                if header_info["sao_y"] is None:
                    header_info["sao_y"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            # 2. Số ký hiệu
            so_match = re.search(r"(S[oố]\s*:\s*[0-9A-Z/a-z_#-]+)", segment, re.IGNORECASE)
            if so_match and idx < 12:
                if header_info["so_ky_hieu"] is None:
                    header_info["so_ky_hieu"] = so_match.group(1).strip()
                    header_end_idx = max(header_end_idx, idx)
                
                # If the segment contains other text before/after "Số: ...", keep processing it
                rem_text = segment.replace(so_match.group(0), "").strip()
                if not rem_text:
                    continue
                segment = rem_text
                segment_upper = segment.upper()
                
            # 3. Ngày tháng
            if re.search(r"ngày\s*\d+\s*tháng\s*\d+\s*năm\s*\d+", segment, re.IGNORECASE) and idx < 12:
                if header_info["ngay_thang"] is None:
                    header_info["ngay_thang"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            # 4. Quốc hiệu & Tiêu ngữ
            if re.search(r"CỘNG\s*H[OÒ]À?\s*X[AÃ]", segment_upper) and idx < 10:
                if header_info["quoc_hieu"] is None:
                    tieu_ngu_match = re.search(r"(Đ[oộ]c\s*l[aậ]p\s*[-–—]\s*T[uự]\s*do\s*[-–—]\s*H[aạ]nh\s*ph[uú]c)", segment, re.IGNORECASE)
                    if tieu_ngu_match:
                        header_info["tieu_ngu"] = tieu_ngu_match.group(1).strip()
                        quoc_hieu_text = segment[:tieu_ngu_match.start()].strip()
                        quoc_hieu_text = re.sub(r"\s*[-–—]\s*$", "", quoc_hieu_text).strip()
                        header_info["quoc_hieu"] = quoc_hieu_text if quoc_hieu_text else "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
                    else:
                        header_info["quoc_hieu"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            if re.search(r"Đ[oộ]c\s*l[aậ]p.*T[uự]\s*do.*H[aạ]nh\s*ph[uú]c", segment, re.IGNORECASE) and idx < 10:
                if header_info["tieu_ngu"] is None:
                    header_info["tieu_ngu"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            # 5. Cơ quan ban hành
            if re.search(r"([UỦ][YỶÝỲỸ]?\s*BAN|^BỘ\s|^SỞ\s|^UBND|^HỘI\s*ĐỒNG)", segment_upper) and idx < 10:
                if header_info["co_quan_line1"] is None:
                    cq_text = re.sub(r"\s*[-–—]\s*$", "", segment).strip()
                    header_info["co_quan_line1"] = cq_text
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            # 6. Tên văn bản
            if re.match(r"^(TỜ\s+TRÌNH|QUYẾT\s+ĐỊNH|CÔNG\s+VĂN|THÔNG\s+BÁO|BÁO\s+CÁO|KẾ\s+HOẠCH|BIÊN\s+BẢN|GIẤY)", segment_upper) and idx < 15:
                if header_info["ten_van_ban"] is None:
                    header_info["ten_van_ban"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            # 7. Trích yếu
            if re.match(r"^[Vv]ề\s+việc", segment) and idx < 18:
                if header_info["trich_yeu"] is None:
                    header_info["trich_yeu"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
            # 8. Kính gửi
            if re.match(r"^[Kk]ính\s+gửi\s*:", segment) and idx < 20:
                if header_info["kinh_gui"] is None:
                    header_info["kinh_gui"] = segment
                    header_end_idx = max(header_end_idx, idx)
                continue
                
    if header_info["co_quan_line1"] and header_info["co_quan_line2"] is None:
        co_text = header_info["co_quan_line1"].replace("**", "").strip()
        split_match = re.match(r"(.*?[UỦ][YỶÝỲỸ]?\s*BAN\s+NH[AÂ]N\s+D[AÂ]N)\s+(THÀNH\s*PH[OỐ].*)", co_text, re.IGNORECASE)
        if split_match:
            header_info["co_quan_line1"] = split_match.group(1).strip()
            header_info["co_quan_line2"] = split_match.group(2).strip()
            
    has_quoc_hieu = header_info["quoc_hieu"] is not None
    has_co_quan = header_info["co_quan_line1"] is not None
    
    if not (has_quoc_hieu and has_co_quan):
        print("[INFO] Không đủ thông tin header NĐ 30. Bỏ qua tái cấu trúc.")
        return
        
    print(f"[INFO] Phát hiện header NĐ 30 (paragraphs 0-{header_end_idx})")
    
    paras_to_remove = []
    for idx in range(header_end_idx + 1):
        paras_to_remove.append(doc.paragraphs[idx])
        
    body_element = doc.element.body
    for para in reversed(paras_to_remove):
        body_element.remove(para._element)
        
    print(f"[INFO] Đã xóa {len(paras_to_remove)} paragraphs header cũ")
    
    first_element = body_element[0] if len(body_element) > 0 else None
    
    # 3a: Bảng 2 cột
    header_table = doc.add_table(rows=2, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set tblBorders="nil" for the whole table so Layer 3 can skip it
    remove_table_borders(header_table)
    
    set_cell_no_border(header_table.cell(0, 0))
    set_cell_no_border(header_table.cell(0, 1))
    set_cell_no_border(header_table.cell(1, 0))
    set_cell_no_border(header_table.cell(1, 1))
    
    set_table_fixed_width(header_table, pct="5000")
    
    left_w = 3402
    right_w = 5670
    apply_table_grid(header_table, [left_w, right_w])
    
    row1 = header_table.rows[0]
    cell_left_1 = row1.cells[0]
    apply_tcW(cell_left_1, left_w)
    p_cq = cell_left_1.paragraphs[0]
    p_cq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cq.paragraph_format.space_before = Pt(0)
    p_cq.paragraph_format.space_after = Pt(0)
    p_cq.paragraph_format.first_line_indent = Cm(0)
    
    co_quan_text = (header_info["co_quan_line1"] or "").replace("**", "").strip()
    if header_info["co_quan_line2"]:
        co_quan_line2 = header_info["co_quan_line2"].replace("**", "").strip()
        make_run(p_cq, co_quan_text, bold=False, size=12)
        p_cq.add_run("\n")
        make_run(p_cq, co_quan_line2, bold=True, size=13)
        add_separator_line(cell_left_1, width_cm=3.5)
    else:
        lines = co_quan_text.split("\n")
        if len(lines) >= 2:
            make_run(p_cq, lines[0].strip(), bold=False, size=12)
            p_cq.add_run("\n")
            make_run(p_cq, lines[1].strip(), bold=True, size=13)
            add_separator_line(cell_left_1, width_cm=3.5)
        else:
            make_run(p_cq, co_quan_text, bold=True, size=13)
            add_separator_line(cell_left_1, width_cm=3.5)
            
    cell_right_1 = row1.cells[1]
    apply_tcW(cell_right_1, right_w)
    p_qh = cell_right_1.paragraphs[0]
    p_qh.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qh.paragraph_format.space_before = Pt(0)
    p_qh.paragraph_format.space_after = Pt(0)
    p_qh.paragraph_format.first_line_indent = Cm(0)
    
    quoc_hieu = (header_info["quoc_hieu"] or "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM").replace("**", "").strip()
    make_run(p_qh, quoc_hieu, bold=True, size=13)
    p_qh.add_run("\n")
    tieu_ngu = (header_info["tieu_ngu"] or "Độc lập - Tự do - Hạnh phúc").replace("**", "").strip()
    make_run(p_qh, tieu_ngu, bold=True, size=13)
    add_separator_line(cell_right_1, width_cm=5.0)
    
    row2 = header_table.rows[1]
    cell_left_2 = row2.cells[0]
    apply_tcW(cell_left_2, left_w)
    p_so = cell_left_2.paragraphs[0]
    p_so.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_so.paragraph_format.space_before = Pt(6)
    p_so.paragraph_format.space_after = Pt(6)
    p_so.paragraph_format.first_line_indent = Cm(0)
    so_text = (header_info["so_ky_hieu"] or "").replace("**", "").strip()
    make_run(p_so, so_text, size=13)
    
    cell_right_2 = row2.cells[1]
    apply_tcW(cell_right_2, right_w)
    p_ng = cell_right_2.paragraphs[0]
    p_ng.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ng.paragraph_format.space_before = Pt(6)
    p_ng.paragraph_format.space_after = Pt(6)
    p_ng.paragraph_format.first_line_indent = Cm(0)
    ngay_text = (header_info["ngay_thang"] or "").replace("**", "").replace("*", "").strip()
    make_run(p_ng, ngay_text, italic=True, size=13)
    
    body_element.remove(header_table._tbl)
    if first_element is not None:
        body_element.insert(body_element.index(first_element), header_table._tbl)
    else:
        body_element.append(header_table._tbl)
        
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(6)
    spacer.paragraph_format.first_line_indent = Cm(0)
    body_element.remove(spacer._element)
    body_element.insert(body_element.index(header_table._tbl) + 1, spacer._element)
    
    insert_after = spacer._element
    
    if header_info["ten_van_ban"]:
        p_ten = doc.add_paragraph()
        p_ten.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ten.paragraph_format.space_before = Pt(12)
        p_ten.paragraph_format.space_after = Pt(6)
        p_ten.paragraph_format.first_line_indent = Cm(0)
        ten_text = header_info["ten_van_ban"].replace("**", "").replace("##", "").strip()
        make_run(p_ten, ten_text, bold=True, size=14)
        body_element.remove(p_ten._element)
        body_element.insert(body_element.index(insert_after) + 1, p_ten._element)
        insert_after = p_ten._element
        
    if header_info["trich_yeu"]:
        p_ty = doc.add_paragraph()
        p_ty.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ty.paragraph_format.space_before = Pt(0)
        p_ty.paragraph_format.space_after = Pt(6)
        p_ty.paragraph_format.first_line_indent = Cm(0)
        ty_text = header_info["trich_yeu"].replace("**", "").strip()
        make_run(p_ty, ty_text, bold=True, size=14)
        body_element.remove(p_ty._element)
        body_element.insert(body_element.index(insert_after) + 1, p_ty._element)
        insert_after = p_ty._element
        
    p_sep_ty = doc.add_paragraph()
    p_sep_ty.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep_ty.paragraph_format.space_before = Pt(2)
    p_sep_ty.paragraph_format.space_after = Pt(12)
    p_sep_ty.paragraph_format.first_line_indent = Cm(0)
    run_sep = p_sep_ty.add_run("\u2500" * int(3.5 * 4))
    run_sep.font.name = body_font
    run_sep.font.size = Pt(8)
    run_sep.font.color.rgb = RGBColor(0, 0, 0)
    body_element.remove(p_sep_ty._element)
    body_element.insert(body_element.index(insert_after) + 1, p_sep_ty._element)
    insert_after = p_sep_ty._element
    
    if header_info["kinh_gui"]:
        p_kg = doc.add_paragraph()
        p_kg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_kg.paragraph_format.space_before = Pt(12)
        p_kg.paragraph_format.space_after = Pt(12)
        p_kg.paragraph_format.first_line_indent = Cm(0)
        kg_text = header_info["kinh_gui"].replace("**", "").strip()
        if kg_text.lower().startswith("kính gửi"):
            kg_text = kg_text[8:].strip(" :")
            make_run(p_kg, f"Kính gửi: {kg_text}", size=14)
        else:
            make_run(p_kg, kg_text, size=14)
        body_element.remove(p_kg._element)
        body_element.insert(body_element.index(insert_after) + 1, p_kg._element)
        
    print("[OK] Header NĐ 30 đã tái cấu trúc")
    doc.save(str(docx_path))

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python 02_structure.py <input_docx>")
        sys.exit(1)
        
    input_docx = Path(sys.argv[1])
    print("[INFO] Bước 2: Tái cấu trúc Header NĐ 30...")
    build_nd30_header(input_docx)
    print("[OK] Xong Layer 2.")
