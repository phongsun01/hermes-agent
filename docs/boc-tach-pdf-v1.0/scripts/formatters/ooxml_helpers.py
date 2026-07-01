import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, Cm, RGBColor

# Function to safely parse XML
def parse_xml(xml_str):
    """Wrapper to parse xml string to oxml element safely."""
    from lxml import etree
    return etree.fromstring(xml_str)

def make_run(para, text, body_font="Times New Roman", bold=False, italic=False, size=13, underline=False):
    """Tạo run với font chuẩn và Complex Script tags cho tiếng Việt."""
    run = para.add_run(text)
    run.font.name = body_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    if underline:
        run.font.underline = True
        
    rpr = run._element.get_or_add_rPr()
    if bold:
        if rpr.find(qn('w:bCs')) is None:
            rpr.append(OxmlElement('w:bCs'))
    if italic:
        if rpr.find(qn('w:iCs')) is None:
            rpr.append(OxmlElement('w:iCs'))
            
    # Set eastAsia/cs
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), body_font)
    rfonts.set(qn('w:hAnsi'), body_font)
    rfonts.set(qn('w:eastAsia'), body_font)
    rfonts.set(qn('w:cs'), body_font)
    return run

def set_cell_no_border(cell):
    """Xóa viền ô."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcPr.append(parse_xml(f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="nil"/>
        <w:left w:val="nil"/>
        <w:bottom w:val="nil"/>
        <w:right w:val="nil"/>
    </w:tcBorders>'''))

def add_separator_line(cell, width_cm=3.5, body_font="Times New Roman"):
    """Thêm đường kẻ ngắn căn giữa dưới dòng chữ."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p_sep = cell.add_paragraph()
    p_sep.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(2)
    p_sep.paragraph_format.first_line_indent = Cm(0)
    
    run = p_sep.add_run("\u2500" * int(width_cm * 4))  # Box drawing horizontal
    run.font.name = body_font
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = False
    return p_sep

def apply_tcW(cell, width_dxa):
    """Gắn width_dxa cố định vào cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')

def remove_table_borders(table):
    """Xóa toàn bộ viền bảng."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(parse_xml(f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="nil"/>
        <w:left w:val="nil"/>
        <w:bottom w:val="nil"/>
        <w:right w:val="nil"/>
        <w:insideH w:val="nil"/>
        <w:insideV w:val="nil"/>
    </w:tblBorders>'''))
    return tblPr

def set_table_fixed_width(table, pct="5000"):
    """Thiết lập w:tblW = pct và w:tblLayout = fixed."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
        
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{pct}" w:type="pct"/>'))
    
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

def apply_table_grid(table, cols_width_dxa):
    """Thiết lập w:tblGrid cho bảng dựa trên mảng cols_width_dxa."""
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in cols_width_dxa:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += '</w:tblGrid>'
    
    tbl.insert(1, parse_xml(grid_xml))  # After tblPr
