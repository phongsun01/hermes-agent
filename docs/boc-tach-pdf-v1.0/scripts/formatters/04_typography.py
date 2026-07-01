import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from ooxml_helpers import parse_xml, nsdecls, qn, OxmlElement

def set_run_font(run, font_name, size_pt, bold=None, italic=None, color=None):
    """Set font cho run. bold/italic=None → giữ nguyên giá trị hiện có."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    
    rpr = run._element.get_or_add_rPr()
    
    if bold is not None:
        run.font.bold = bold
        if bold:
            if rpr.find(qn('w:bCs')) is None:
                rpr.append(OxmlElement('w:bCs'))
        else:
            for bcs in rpr.findall(qn('w:bCs')):
                rpr.remove(bcs)
                
    if italic is not None:
        run.font.italic = italic
        if italic:
            if rpr.find(qn('w:iCs')) is None:
                rpr.append(OxmlElement('w:iCs'))
        else:
            for ics in rpr.findall(qn('w:iCs')):
                rpr.remove(ics)
                
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    # Đảm bảo eastAsia/cs
    rf = rpr.find(qn('w:rFonts'))
    if rf is not None:
        for attr in ('w:eastAsia', 'w:cs'):
            rf.set(qn(attr), font_name)

def format_typography(docx_path, format_spec):
    """
    Định dạng cấp hạt (Inline Level):
    - Dọn rác style Normal (ví dụ: <w:b w:val="0"/>).
    - Duyệt qua từng Run, ép font chữ, size, và color đen trắng cho NĐ 30.
    - Ép thẻ Complex Script cho Bold/Italic tiếng Việt.
    """
    doc = Document(str(docx_path))
    
    body_font = format_spec.get("body_font", "Times New Roman")
    body_size = format_spec.get("body_font_size", 13.0)
    doc_type = format_spec.get("doc_type", "van_ban_dai")
    
    # 1. Dọn rác style Normal
    if 'Normal' in doc.styles:
        normal_style = doc.styles['Normal']
        style_rPr = normal_style.element.find(qn('w:rPr'))
        if style_rPr is not None:
            # Xóa b val=0 để toggle in đậm hoạt động
            for b_elem in style_rPr.findall(qn('w:b')):
                style_rPr.remove(b_elem)
                print("[FIX] Đã xóa <w:b w:val='0'/> khỏi Normal style")
            
            # Xóa justify style (nếu có, để direct formatting quyết định)
            # Normal style paragraph properties
            style_pPr = normal_style.element.find(qn('w:pPr'))
            if style_pPr is not None:
                for jc_elem in style_pPr.findall(qn('w:jc')):
                    style_pPr.remove(jc_elem)
                    print("[FIX] Đã xóa <w:jc/> khỏi Normal style")
                    
    # 2. Xử lý từng paragraph runs
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        
        # Nếu là Heading trong NĐ 30: ép bold đen
        is_heading_nd30 = doc_type == "hanh_chinh_nd30" and style_name.startswith('Heading')
        
        for run in para.runs:
            if is_heading_nd30:
                set_run_font(run, body_font, body_size, bold=True)
            elif doc_type == "hanh_chinh_nd30":
                # NĐ 30 body text: force font, size, đen, GIỮ NGUYÊN trạng thái bold/italic
                set_run_font(run, body_font, body_size)
            else:
                # Văn bản dài: chỉ ép lại font và color đen nếu cần
                set_run_font(run, body_font, body_size)

    # 3. Xử lý runs trong Table cells
    if doc_type == "hanh_chinh_nd30":
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            # Headers bảng
                            if i == 0:
                                set_run_font(run, body_font, 11, bold=True)
                            else:
                                set_run_font(run, body_font, 11)

    doc.save(str(docx_path))

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python 04_typography.py <input_docx> <format_spec.json>")
        sys.exit(1)
        
    input_docx = Path(sys.argv[1])
    format_spec_path = Path(sys.argv[2])
    
    with open(format_spec_path, "r", encoding="utf-8") as f:
        format_spec = json.load(f)
        
    print("[INFO] Bước 4: Định dạng Typography (Font, Bold/Italic tags)...")
    format_typography(input_docx, format_spec)
    print("[OK] Xong Layer 4.")
