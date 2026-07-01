import os
import sys
import json
import re
import shutil
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt

def enrich_md_with_images(merged_md_path, image_map_path, images_dir):
    """
    Đọc MERGED.md, chèn ảnh minh họa (nếu có image_map),
    và LUÔN chuyển <!-- page: N --> thành [PAGE_MARKER_N].
    """
    with open(merged_md_path, "r", encoding="utf-8") as f:
        md_text = f.read()
    
    # Chèn ảnh minh họa (nếu có image_map)
    if image_map_path.exists():
        with open(image_map_path, "r", encoding="utf-8") as f:
            image_map = json.load(f)
        
        def replacer(match):
            img_name = match.group(1).strip()
            if img_name in image_map:
                rel_path = f"../01.input/{image_map[img_name]}"
                return f"![{img_name}]({rel_path})"
            return match.group(0)
        
        md_text = re.sub(r"\[Hình minh họa: (.*?)\]", replacer, md_text)
    
    # LUÔN chuyển <!-- page: N --> thành [PAGE_MARKER_N] để Pandoc không xóa mất
    def page_marker_replacer(match):
        page_num = match.group(1).strip()
        return f"[PAGE_MARKER_{page_num}]"
        
    enriched_text = re.sub(r"<!--\s*page:\s*(\d+)\s*-->", page_marker_replacer, md_text)
    
    enriched_path = merged_md_path.parent / "ENRICHED.md"
    with open(enriched_path, "w", encoding="utf-8") as f:
        f.write(enriched_text)
        
    return enriched_path

def export_with_pandoc(enriched_md, reference_docx, output_docx, images_dir):
    """Sử dụng pypandoc-binary để chuyển đổi sang DOCX."""
    try:
        import pypandoc
        
        args = [
            f'--reference-doc={reference_docx}',
            '--wrap=none',
            '+RTS', '-V0', '-RTS'
        ]
        
        pypandoc.convert_file(
            str(enriched_md),
            'docx',
            outputfile=str(output_docx),
            extra_args=args
        )
        return True
    except ImportError:
        print("[WARN] Chưa cài đặt pypandoc-binary (pip install pypandoc-binary).")
        return False
    except Exception as e:
        print(f"[WARN] Pandoc lỗi: {e}")
        return False

def export_fallback_docx(enriched_md, format_spec, output_docx):
    """Fallback xuất DOCX bằng python-docx (không dùng reference.docx)."""
    print("[INFO] Sử dụng fallback python-docx (basic formatting)...")
    doc = Document()
    
    with open(enriched_md, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.startswith("<!-- page:"):
                doc.add_paragraph(f"[{line.strip('<!-- >')}]")
                continue
            
            p = doc.add_paragraph()
            if line.startswith("## "):
                run = p.add_run(line[3:])
                run.font.size = Pt(16)
                run.font.bold = True
            elif line.startswith("### "):
                run = p.add_run(line[4:])
                run.font.size = Pt(14)
                run.font.bold = True
            else:
                p.add_run(line)
                
    doc.save(str(output_docx))
    return True

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python 00_pandoc.py <processing_dir> <output_docx>")
        sys.exit(1)
        
    processing_dir = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])
    
    process_dir = processing_dir / "02.process"
    output_dir = processing_dir / "03.output"
    input_dir = processing_dir / "01.input"
    merged_md = process_dir / "MERGED.md"
    reference_docx = process_dir / "reference.docx"
    image_map_path = process_dir / "image_map.json"
    format_spec_path = process_dir / "format_spec.json"
    
    if not merged_md.exists():
        print("[FAIL] Không tìm thấy MERGED.md")
        sys.exit(1)
        
    with open(format_spec_path, "r", encoding="utf-8") as f:
        format_spec = json.load(f)
        
    print("[INFO] Bước 0: Chèn ảnh minh họa vào Markdown...")
    enriched = enrich_md_with_images(merged_md, image_map_path, input_dir)
    
    print("[INFO] Bước 0: Xuất DOCX bằng Pandoc + reference template...")
    success = False
    if reference_docx.exists():
        success = export_with_pandoc(enriched, reference_docx, output_docx, input_dir)
        
    if not success:
        print("[INFO] Pandoc không khả dụng. Chuyển sang fallback...")
        success = export_fallback_docx(enriched, format_spec, output_docx)
        
    if success:
        print(f"[OK] Đã xuất DOCX thô thành công: {output_docx.name}")
    else:
        print("[FAIL] Xuất DOCX thất bại")
        sys.exit(1)
