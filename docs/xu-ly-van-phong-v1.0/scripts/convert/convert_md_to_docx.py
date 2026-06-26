"""
Script chuyển đổi MD sang DOCX theo format template
Sử dụng python-docx để tạo DOCX với format chuẩn văn bản hành chính
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import yaml

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # Set font for Asian characters
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph_with_format(doc, text, alignment='left', font_size=12, bold=False, italic=False, spacing_after=6):
    """Add a paragraph with specific formatting"""
    para = doc.add_paragraph()
    
    # Set alignment
    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set spacing
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = 1.15
    
    # Add text with formatting
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold, italic=italic)
    
    return para

def parse_md_line(line):
    """Parse markdown formatting from a line"""
    # Remove ** for bold markers
    text = line.strip()
    is_bold = text.startswith('**') and '**' in text[2:]
    is_italic = text.startswith('*') and not text.startswith('**')
    
    # Clean markdown syntax
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Remove italic markers
    
    return text, is_bold, is_italic

def add_mixed_paragraph(doc, text, alignment='justify', base_size=12, spacing_after=6):
    """Add paragraph with mixed bold/normal text"""
    para = doc.add_paragraph()
    
    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = 1.15
    
    # Parse bold sections **text** and italic *text*
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, font_size=base_size, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            set_run_font(run, font_size=base_size, italic=True)
        else:
            run = para.add_run(part)
            set_run_font(run, font_size=base_size)
    
    return para

def add_nd30_header(doc, org_parent, org_name, so_ky_hieu, date_str):
    """Add ND 30 standard 2-column header"""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8.5)
    
    # Row 0: Org vs Quốc hiệu
    cell_00 = table.cell(0, 0)
    p00 = cell_00.paragraphs[0]
    p00.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p00.paragraph_format.space_after = Pt(0)
    if org_parent:
        r_parent = p00.add_run(org_parent.upper() + '\n')
        set_run_font(r_parent, font_size=12)
    r_org = p00.add_run(org_name.upper())
    set_run_font(r_org, font_size=13, bold=True)
    
    cell_01 = table.cell(0, 1)
    p01 = cell_01.paragraphs[0]
    p01.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p01.paragraph_format.space_after = Pt(0)
    r_qh = p01.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
    set_run_font(r_qh, font_size=13, bold=True)
    r_td = p01.add_run('Độc lập - Tự do - Hạnh phúc')
    set_run_font(r_td, font_size=14, bold=True)
    
    # Row 1: Số ký hiệu vs Ngày tháng
    cell_10 = table.cell(1, 0)
    p10 = cell_10.paragraphs[0]
    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_so = p10.add_run(so_ky_hieu)
    set_run_font(r_so, font_size=13)
    
    cell_11 = table.cell(1, 1)
    p11 = cell_11.paragraphs[0]
    p11.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_date = p11.add_run(date_str)
    set_run_font(r_date, font_size=13, italic=True)
    
    doc.add_paragraph()

def convert_md_to_docx(md_file, output_file):
    """Convert MD file to DOCX with proper formatting"""
    
    # Read MD content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse YAML frontmatter if exists
    frontmatter = {}
    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            try:
                frontmatter = yaml.safe_load(parts[1])
                lines = parts[2].splitlines()
            except Exception as e:
                print(f"Error parsing YAML frontmatter: {e}")
                lines = content.splitlines()
        else:
            lines = content.splitlines()
    else:
        lines = content.splitlines()
    
    # Create new document
    doc = Document()
    
    # Set up page margins (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2)
    
    # Process ND30 header if requested
    if frontmatter.get('nd30_header'):
        add_nd30_header(
            doc,
            frontmatter.get('org_parent', ''),
            frontmatter.get('org_name', ''),
            frontmatter.get('so_ky_hieu', 'Số:      /'),
            frontmatter.get('date', 'Ngày      tháng      năm 202...')
        )
    
    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Header: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
        if 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM' in line:
            add_paragraph_with_format(doc, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', 
                                     alignment='center', font_size=14, bold=True)
            i += 1
            continue
        
        # Độc lập - Tự do - Hạnh phúc
        if 'Độc lập' in line and 'Tự do' in line:
            add_paragraph_with_format(doc, 'Độc lập - Tự do - Hạnh phúc', 
                                     alignment='center', font_size=14, bold=True)
            i += 1
            continue
        
        # Date line (right aligned, italic)
        if line.startswith('Bình Dương, ngày'):
            add_paragraph_with_format(doc, line, alignment='right', font_size=12, italic=True)
            i += 1
            continue
        
        # Main title: ĐƠN TỐ CÁO TỘI PHẠM
        if 'ĐƠN TỐ CÁO TỘI PHẠM' in line or 'ĐƠN TỐ GIÁC TỘI PHẠM' in line:
            clean_title = line.replace('**', '')
            add_paragraph_with_format(doc, clean_title, alignment='center', font_size=14, bold=True, spacing_after=0)
            i += 1
            continue
        
        # Subtitle (italic, centered)
        if line.startswith('*Về hành vi'):
            clean_text = line.replace('*', '')
            add_paragraph_with_format(doc, clean_text, alignment='center', font_size=12, italic=True)
            i += 1
            continue
        
        # Section headers (I., II., III., etc.)
        if re.match(r'^\*\*[IVX]+\.', line):
            clean_text = line.replace('**', '')
            add_paragraph_with_format(doc, clean_text, alignment='left', font_size=12, bold=True, spacing_after=6)
            i += 1
            continue
        
        # Sub-headers within sections (1., 2., 3.)
        if re.match(r'^\*\*\d+\.', line):
            clean_text = line.replace('**', '')
            add_paragraph_with_format(doc, clean_text, alignment='left', font_size=12, bold=True, spacing_after=6)
            i += 1
            continue
        
        # Kính gửi
        if line.startswith('**Kính gửi:**'):
            add_paragraph_with_format(doc, 'Kính gửi:', alignment='left', font_size=12, bold=True, spacing_after=0)
            i += 1
            continue
        
        # List items (- or numbered)
        if line.startswith('- ') or line.startswith('+ '):
            clean_text = line[2:].strip()
            para = add_mixed_paragraph(doc, '• ' + clean_text, alignment='justify', base_size=12)
            para.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue
        
        # Numbered list items (1., 2., etc. at start)
        if re.match(r'^\d+\.\s', line):
            add_mixed_paragraph(doc, line, alignment='justify', base_size=12)
            i += 1
            continue
        
        # Sub-list items with indent
        if line.startswith('   -') or line.startswith('   +'):
            clean_text = line.strip()[2:].strip()
            para = add_mixed_paragraph(doc, '  ○ ' + clean_text, alignment='justify', base_size=12)
            para.paragraph_format.left_indent = Cm(1)
            i += 1
            continue
        
        # Signature section
        if 'CÔNG TY CỔ PHẦN NỀN TẢNG' in line:
            add_paragraph_with_format(doc, '', alignment='left', spacing_after=24)  # Space before
            add_paragraph_with_format(doc, 'CÔNG TY CỔ PHẦN NỀN TẢNG', alignment='right', font_size=12, bold=True, spacing_after=0)
            i += 1
            continue
        
        if 'TÀI TRỢ CHUỖI CUNG ỨNG SCP' in line:
            add_paragraph_with_format(doc, 'TÀI TRỢ CHUỖI CUNG ỨNG SCP', alignment='right', font_size=12, bold=True)
            i += 1
            continue
        
        if 'Người đại diện pháp luật' in line and 'Ký' not in line:
            add_paragraph_with_format(doc, 'Người đại diện pháp luật', alignment='right', font_size=12, spacing_after=0)
            i += 1
            continue
        
        if '(Ký tên, đóng dấu)' in line:
            add_paragraph_with_format(doc, '(Ký tên, đóng dấu)', alignment='right', font_size=12, italic=True, spacing_after=48)
            i += 1
            continue
        
        if 'Lê Minh Đức' in line:
            add_paragraph_with_format(doc, 'Lê Minh Đức', alignment='right', font_size=12, bold=True, spacing_after=0)
            i += 1
            continue
        
        if 'Tổng giám đốc' in line and 'Chức vụ' not in line:
            add_paragraph_with_format(doc, 'Tổng giám đốc', alignment='right', font_size=12)
            i += 1
            continue
        
        # Note at bottom
        if line.startswith('*Lưu ý:'):
            add_paragraph_with_format(doc, '', alignment='left', spacing_after=12)
            clean_text = line.replace('*', '')
            add_paragraph_with_format(doc, clean_text, alignment='left', font_size=11, italic=True, spacing_after=0)
            i += 1
            continue
        
        # Regular paragraph with mixed formatting
        add_mixed_paragraph(doc, line, alignment='justify', base_size=12)
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"Created: {output_file}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
