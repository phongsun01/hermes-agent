"""
Professional Vietnamese Textbook DOCX Formatter v2
Specs from user screenshot:
  - Justified, First line 1cm, Before 6pt, After 6pt, Line spacing 1.15
  - Cover page: centered mid-page, bold uppercase large
  - Image captions: centered
  - Tables: header bold+centered, body left-aligned, auto-fit width
  - Images: centered, no indent
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def format_docx(input_path, output_path=None):
    if output_path is None:
        output_path = input_path

    doc = Document(input_path)

    # ═══════════ 1. PAGE SETUP ═══════════
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    # ═══════════ 1b. OVERRIDE BULLET CHARACTERS ═══════════
    # Pandoc stores bullet chars in numbering.xml. Override level 0 → dash, level 1 → circle
    numbering_part = None
    for rel in doc.part.rels.values():
        if 'numbering' in rel.reltype:
            numbering_part = rel.target_part
            break
    if numbering_part is not None:
        num_xml = numbering_part.element
        for abstractNum in num_xml.findall(qn('w:abstractNum')):
            for lvl in abstractNum.findall(qn('w:lvl')):
                ilvl_val = lvl.get(qn('w:ilvl'), '0')
                numFmt = lvl.find(qn('w:numFmt'))
                # Only modify bullet lists (not numbered)
                if numFmt is not None and numFmt.get(qn('w:val')) == 'bullet':
                    lvlText = lvl.find(qn('w:lvlText'))
                    # Set font to Arial (supports all Unicode bullet chars)
                    rPr = lvl.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                        lvl.append(rPr)
                    for old_rf in rPr.findall(qn('w:rFonts')):
                        rPr.remove(old_rf)
                    rPr.append(parse_xml(
                        f'<w:rFonts {nsdecls("w")} w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
                    ))
                    if lvlText is not None:
                        if ilvl_val == '0':
                            lvlText.set(qn('w:val'), '\u2013')  # en-dash –
                        elif ilvl_val == '1':
                            lvlText.set(qn('w:val'), '\u25CF')  # filled circle ●

    # ═══════════ 2. STYLE DEFINITIONS ═══════════
    styles = doc.styles

    def set_font_all(style, font_name='Times New Roman'):
        """Set font on both run properties and eastAsia/cs"""
        style.font.name = font_name
        rpr = style.element.find(qn('w:rPr'))
        if rpr is None:
            rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            style.element.append(rpr)
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rpr.insert(0, rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rfonts.set(qn(attr), font_name)

    # --- Normal (Body Text) - MATCHING USER SCREENSHOT ---
    normal = styles['Normal']
    set_font_all(normal)
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15

    # --- Heading 1 ---
    h1 = styles['Heading 1']
    set_font_all(h1)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    h1f = h1.paragraph_format
    h1f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1f.space_before = Pt(24)
    h1f.space_after = Pt(12)
    h1f.first_line_indent = Cm(0)
    h1f.keep_with_next = True
    h1f.page_break_before = False  # Don't force page break on style (handle per-paragraph)

    # --- Heading 2 ---
    h2 = styles['Heading 2']
    set_font_all(h2)
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    h2f = h2.paragraph_format
    h2f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2f.space_before = Pt(18)
    h2f.space_after = Pt(6)
    h2f.first_line_indent = Cm(0)
    h2f.keep_with_next = True
    h2f.page_break_before = False

    # --- Heading 3 ---
    h3 = styles['Heading 3']
    set_font_all(h3)
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = RGBColor(0x00, 0x69, 0x7A)  # Dark teal
    h3f = h3.paragraph_format
    h3f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3f.space_before = Pt(6)
    h3f.space_after = Pt(3)
    h3f.first_line_indent = Cm(0)
    h3f.keep_with_next = True

    # --- Heading 4 ---
    if 'Heading 4' in [s.name for s in styles]:
        h4 = styles['Heading 4']
    else:
        h4 = styles.add_style('Heading 4', 1)  # 1 = paragraph
    set_font_all(h4)
    h4.font.size = Pt(13)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0xBF, 0x36, 0x0C)  # Orange-brown
    h4f = h4.paragraph_format
    h4f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4f.space_before = Pt(10)
    h4f.space_after = Pt(4)
    h4f.first_line_indent = Cm(0)
    h4f.keep_with_next = True

    # ═══════════ 3. PARAGRAPHS ═══════════
    def add_shading(element, color):
        """Add background shading to a paragraph element."""
        pPr = element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            element.insert(0, pPr)
        for old in pPr.findall(qn('w:shd')):
            pPr.remove(old)
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        ))

    def add_border_box(element, color='CCCCCC'):
        """Add a border box around a paragraph."""
        pPr = element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            element.insert(0, pPr)
        for old in pPr.findall(qn('w:pBdr')):
            pPr.remove(old)
        pPr.append(parse_xml(f'''<w:pBdr {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="4" w:color="{color}"/>
            <w:left w:val="single" w:sz="4" w:space="8" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{color}"/>
            <w:right w:val="single" w:sz="4" w:space="8" w:color="{color}"/>
        </w:pBdr>'''))

    def add_bottom_border(element, color='1B5E20', size='6'):
        """Add a bottom border line under a paragraph (for H2 underline)."""
        pPr = element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            element.insert(0, pPr)
        for old in pPr.findall(qn('w:pBdr')):
            pPr.remove(old)
        pPr.append(parse_xml(f'''<w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="{size}" w:space="3" w:color="{color}"/>
        </w:pBdr>'''))

    def add_hr_border(element):
        """Style empty paragraph as a horizontal rule."""
        pPr = element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            element.insert(0, pPr)
        for old in pPr.findall(qn('w:pBdr')):
            pPr.remove(old)
        pPr.append(parse_xml(f'''<w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>
        </w:pBdr>'''))

    # Detect if this is a proposal/report (not a book with chapters)
    has_chapters = any(p.text.strip().startswith('CHƯƠNG') for p in doc.paragraphs)

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        text = para.text.strip()

        # Force font on all runs (default)
        for run in para.runs:
            if not run.font.name:
                run.font.name = 'Times New Roman'
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                rf = rpr.find(qn('w:rFonts'))
                if rf is not None:
                    for attr in ('w:eastAsia', 'w:cs'):
                        rf.set(qn(attr), 'Times New Roman')

        # --- CODE BLOCKS: gray background + monospace font ---
        if style_name in ('Source Code', 'Verbatim Char') or \
           style_name.startswith('Source') or \
           'code' in style_name.lower():
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.0
            add_shading(para._element, 'F0F4F8')  # Light blue-gray
            add_border_box(para._element, 'D0D8E0')
            code_font = 'Consolas'
            for run in para.runs:
                run.font.name = code_font
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x33, 0x3B)
                rpr = run._element.find(qn('w:rPr'))
                if rpr is not None:
                    rf = rpr.find(qn('w:rFonts'))
                    if rf is not None:
                        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                            rf.set(qn(attr), code_font)
            continue

        # --- HEADINGS ---
        if style_name.startswith('Heading'):
            para.paragraph_format.first_line_indent = Cm(0)
            # H1: center for title, left for chapter sections
            if style_name == 'Heading 1':
                if has_chapters:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Page break for non-first H1 in chapter books
                    para.paragraph_format.page_break_before = (idx > 0)
                else:
                    # Proposal/report: center the main title, no page break for first
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.page_break_before = False
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.space_after = Pt(12)
            elif style_name == 'Heading 2':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(18)
                para.paragraph_format.space_after = Pt(6)
                # Add subtle bottom border for H2 visual separation
                add_bottom_border(para._element, '1B5E20', '4')
            elif style_name == 'Heading 3':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(4)
            elif style_name == 'Heading 4':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(4)
            continue

        # --- BLOCKQUOTES: Pandoc Block Text style → blue left border + light bg ---
        if style_name in ('Block Text', 'Quote', 'Intense Quote') or \
           style_name.startswith('Block'):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(0.3)
            para.paragraph_format.right_indent = Cm(0.3)
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(8)
            add_shading(para._element, 'E8F4FD')  # Light blue
            # Add thick blue left border only
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                para._element.insert(0, pPr)
            for old in pPr.findall(qn('w:pBdr')):
                pPr.remove(old)
            pPr.append(parse_xml(f'''<w:pBdr {nsdecls("w")}>
                <w:left w:val="single" w:sz="24" w:space="8" w:color="1976D2"/>
            </w:pBdr>'''))
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x4D, 0x6E)
                run.font.italic = True
            continue

        # --- INLINE CODE: highlight runs with Verbatim Char style ---
        for run in para.runs:
            run_style = run.style.name if run.style else ''
            if run_style in ('Verbatim Char', 'Source Code Char') or \
               'code' in run_style.lower():
                run.font.name = 'Consolas'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)  # Pink-red
                rpr = run._element.find(qn('w:rPr'))
                if rpr is None:
                    rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                    run._element.insert(0, rpr)
                for old in rpr.findall(qn('w:shd')):
                    rpr.remove(old)
                rpr.append(parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="FFF0F5" w:val="clear"/>'
                ))
                rf = rpr.find(qn('w:rFonts'))
                if rf is not None:
                    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                        rf.set(qn(attr), 'Consolas')

        # Check for images
        p_elem = para._element
        has_drawing = bool(p_elem.findall('.//' + qn('w:drawing')))

        if has_drawing:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
        elif not text:
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            # Check if preceded by or followed by headings (acts as section separator)
            prev_is_heading = (idx > 0 and doc.paragraphs[idx-1].style.name.startswith('Heading'))
            next_is_heading = (idx < len(doc.paragraphs)-1 and doc.paragraphs[idx+1].style.name.startswith('Heading'))
            if not prev_is_heading and not next_is_heading:
                add_hr_border(para._element)
        elif style_name.startswith('List') or style_name in ('Compact', 'List Paragraph'):
            # Bullet / numbered lists: justify, proper indent
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            # Detect list level
            pPr = para._element.find(qn('w:pPr'))
            ilvl = 0
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    ilvlElem = numPr.find(qn('w:ilvl'))
                    if ilvlElem is not None:
                        ilvl = int(ilvlElem.get(qn('w:val'), '0'))
            if ilvl == 0:
                # Level 1: tight indent
                para.paragraph_format.left_indent = Cm(0.8)
                # Override tab to narrow space
                if pPr is not None:
                    for old_ind in pPr.findall(qn('w:ind')):
                        pPr.remove(old_ind)
                    pPr.append(parse_xml(
                        f'<w:ind {nsdecls("w")} w:left="454" w:hanging="227"/>'
                    ))
            else:
                # Level 2+: italic, slightly more indent
                para.paragraph_format.left_indent = Cm(1.4)
                for run in para.runs:
                    run.font.italic = True
                if pPr is not None:
                    for old_ind in pPr.findall(qn('w:ind')):
                        pPr.remove(old_ind)
                    pPr.append(parse_xml(
                        f'<w:ind {nsdecls("w")} w:left="794" w:hanging="227"/>'
                    ))
        else:
            # Body text: enforce spacing at paragraph level
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Metadata lines (bold label: value at doc start) should not indent
            is_metadata = (style_name == 'First Paragraph' and idx < 10 and
                          any(run.font.bold for run in para.runs if run.text.strip()))
            if is_metadata:
                para.paragraph_format.first_line_indent = Cm(0)
            else:
                para.paragraph_format.first_line_indent = Cm(1.0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)

    # ═══════════ 3b. EMPHASIS KEYWORDS ═══════════
    # Highlight strong/bold text containing emphasis keywords
    emphasis_red = ['tuyệt đối không', 'không được', 'nghiêm cấm', 'cấm',
                    'sai lầm', 'thất bại', 'rủi ro', 'nguy hiểm', 'cảnh báo']
    emphasis_green = ['nên', 'khuyến nghị', 'best practice', 'hiệu quả',
                      'tối ưu', 'quan trọng', 'cần thiết', 'bắt buộc']
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.bold:
                continue
            txt_lower = run.text.lower().strip()
            if not txt_lower:
                continue
            # Red emphasis: negative/warning keywords
            if any(kw in txt_lower for kw in emphasis_red):
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)  # Red
                rpr = run._element.find(qn('w:rPr'))
                if rpr is None:
                    rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                    run._element.insert(0, rpr)
                for old in rpr.findall(qn('w:shd')):
                    rpr.remove(old)
                rpr.append(parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="FFF3F3" w:val="clear"/>'
                ))
            # Green emphasis: positive/recommendation keywords
            elif any(kw in txt_lower for kw in emphasis_green):
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)  # Green
                rpr = run._element.find(qn('w:rPr'))
                if rpr is None:
                    rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                    run._element.insert(0, rpr)
                for old in rpr.findall(qn('w:shd')):
                    rpr.remove(old)
                rpr.append(parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F1F8E9" w:val="clear"/>'
                ))

    # ═══════════ 4. COVER PAGES (runs AFTER paragraph loop) ═══════════
    # Find ALL chapter cover pages (multi-chapter support)
    # Each cover has: "CHƯƠNG X" heading + chapter title on next line
    cover_groups = []  # list of (chương_idx, title_idx)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('CHƯƠNG'):
            title_idx = i + 1 if i + 1 < len(doc.paragraphs) else None
            cover_groups.append((i, title_idx))

    for group_idx, (ch_idx, title_idx) in enumerate(cover_groups):
        # Format the "CHƯƠNG X" paragraph
        para = doc.paragraphs[ch_idx]
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        # Add top spacing to push content toward vertical center
        para.paragraph_format.space_before = Pt(200)
        para.paragraph_format.space_after = Pt(24)
        # First cover: no page break (it's at doc start)
        # Other covers: page break before the CHƯƠNG heading
        para.paragraph_format.page_break_before = (group_idx > 0)
        for run in para.runs:
            run.font.size = Pt(32)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
            run.font.name = 'Times New Roman'

        # Format the chapter title paragraph
        if title_idx is not None:
            title_para = doc.paragraphs[title_idx]
            title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.first_line_indent = Cm(0)
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(200)
            title_para.paragraph_format.page_break_before = False
            for run in title_para.runs:
                run.font.size = Pt(28)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
                run.font.name = 'Times New Roman'

    # ═══════════ 5. TABLES ═══════════
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl = table._tbl

        # Set table to auto-fit contents within page width
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)

        # Width = 100% of page
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblPr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
        ))

        # Auto-fit layout
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblPr.append(parse_xml(
            f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>'
        ))

        # Borders: clean gray
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)
        tblPr.append(parse_xml(f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>
        </w:tblBorders>'''))

        # Cell padding
        for old in tblPr.findall(qn('w:tblCellMar')):
            tblPr.remove(old)
        tblPr.append(parse_xml(f'''<w:tblCellMar {nsdecls("w")}>
            <w:top w:w="60" w:type="dxa"/>
            <w:left w:w="100" w:type="dxa"/>
            <w:bottom w:w="60" w:type="dxa"/>
            <w:right w:w="100" w:type="dxa"/>
        </w:tblCellMar>'''))

        # Format rows
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                # Vertical align middle
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                    tc.insert(0, tcPr)

                for para in cell.paragraphs:
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.space_before = Pt(3)
                    para.paragraph_format.space_after = Pt(3)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

                    if i == 0:
                        # HEADER ROW: bold, centered
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        # BODY ROWS: left-aligned
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Header row: blue background
                if i == 0:
                    for old in tcPr.findall(qn('w:shd')):
                        tcPr.remove(old)
                    tcPr.append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="1565C0" w:val="clear"/>'
                    ))
                else:
                    # Alternating row colors for readability
                    bg = 'F5F5F5' if i % 2 == 1 else 'FFFFFF'
                    for old in tcPr.findall(qn('w:shd')):
                        tcPr.remove(old)
                    tcPr.append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'
                    ))

    # ═══════════ 6. CENTER IMAGES VIA XML ═══════════
    for para in doc.paragraphs:
        p_elem = para._element
        if p_elem.findall('.//' + qn('w:drawing')):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                p_elem.insert(0, pPr)
            # Center
            for old in pPr.findall(qn('w:jc')):
                pPr.remove(old)
            pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
            # No indent
            for old in pPr.findall(qn('w:ind')):
                pPr.remove(old)
            pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="0"/>'))

    # ═══════════ SAVE ═══════════
    doc.save(output_path)
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    img_count = sum(1 for p in doc.paragraphs
                    for _ in p._element.findall('.//' + qn('w:drawing')))
    print(f"OK — {output_path}")
    print(f"   {para_count} paragraphs | {table_count} tables | {img_count} images")


if __name__ == '__main__':
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else inp
    format_docx(inp, out)
